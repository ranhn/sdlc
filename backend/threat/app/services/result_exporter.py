"""威胁建模结果的文档导出。

把一次已保存的建模结果（含完整 Threat Dragon 模型）渲染为人类可读的
Markdown 报告 / Threat Dragon 标准 JSON / CSV 威胁清单，便于用户下载 /
存档 / 分享 / 二次编辑。
"""

from __future__ import annotations

import csv
import io
import json
from datetime import datetime
from typing import Any

# severity 排序权重（用于报告内排序）
SEV_ORDER = {
    "Critical": 0,
    "High": 1,
    "Medium": 2,
    "Low": 3,
    "Unknown": 9,
}


def _fmt_time(epoch: float) -> str:
    try:
        return datetime.fromtimestamp(epoch).strftime("%Y-%m-%d %H:%M:%S")
    except (ValueError, OSError, TypeError):
        return "-"


def _sev(key: str) -> str:
    return SEV_ORDER.get(key, SEV_ORDER["Unknown"])


def render_result_markdown(record: dict[str, Any]) -> str:
    """把一条结果记录渲染为 Markdown 报告文本。"""
    model = record.get("model") or {}
    summary = model.get("summary") or {}
    stats = record.get("stats") or {}
    methodology = record.get("methodology") or "STRIDE"
    created = _fmt_time(record.get("created_at", 0))

    lines: list[str] = []
    lines.append(f"# 威胁建模报告：{record.get('title', '未命名')}")
    lines.append("")
    lines.append(f"- **创建时间**：{created}")
    lines.append(f"- **方法论**：{methodology}")
    lines.append(f"- **模型标题**：{summary.get('title', '-')}")
    lines.append(f"- **模型描述**：{summary.get('description', '-') or '-'}")
    lines.append("")

    # 统计概览
    lines.append("## 1、统计概览")
    lines.append("")
    lines.append(f"- 组件数：{stats.get('componentCount', '-')}")
    lines.append(f"- 数据流数：{stats.get('flowCount', '-')}")
    lines.append(f"- 威胁总数：{stats.get('threatCount', '-')}")
    by_sev = stats.get("threatCountBySeverity") or {}
    if by_sev:
        sev_bits = [f"{k}: {v}" for k, v in sorted(
            by_sev.items(), key=lambda kv: _sev(kv[0])
        )]
        lines.append(f"- 威胁按等级：{'，'.join(sev_bits)}")
    industry = stats.get("industry")
    if industry:
        from .ai_knowledge import get_industry_template

        tmpl = get_industry_template(industry)
        lines.append(f"- 行业场景：{tmpl['label'] if tmpl else industry}")
    lines.append("")

    # 度量指标（STRIDE-AI 下输出覆盖度/风险收敛/DREAD/合规）
    metrics = stats.get("metrics") or {}
    if metrics:
        lines.append("## 1.2、度量指标")
        lines.append("")
        lines.append(f"- 元素覆盖度：{_pct(metrics.get('coverageRate'))}（{metrics.get('modeledElements', 0)}/{metrics.get('totalElements', 0)}）")
        lines.append(f"- 高风险威胁收敛率：{_pct(metrics.get('riskConvergence'))}")
        dread_avg = metrics.get("dreadAverage")
        if dread_avg:
            dread_label = {
                "damage": "危害", "reproducibility": "可重复性",
                "exploitability": "可利用性", "affectedUsers": "受影响面",
                "discoverability": "可发现性",
            }
            bits = [f"{dread_label.get(k, k)}:{v}" for k, v in dread_avg.items() if k != "total"]
            lines.append(f"- DREAD 均值：{' / '.join(bits)}（综合 {dread_avg.get('total', 0)}）")
        llm_cov = metrics.get("owaspLlmCoverRate")
        if llm_cov is not None:
            covered = metrics.get("owaspLlmCovered") or []
            lines.append(f"- OWASP Top10 for LLM 覆盖：{_pct(llm_cov)}（{', '.join(covered) if covered else '无'}）")
        compliance = metrics.get("compliance")
        if compliance:
            lines.append("- 合规映射：")
            for item in compliance:
                mark = "✓" if item.get("covered") else "✗"
                lines.append(f"  - {mark} {item['code']} {item['label']}")
        lines.append("")

    # 威胁明细
    threats = _collect_threats(model)
    lines.append(f"## 2、威胁明细（共 {len(threats)} 条）")
    lines.append("")

    if not threats:
        lines.append("_未识别到威胁。_")
    else:
        # 按严重度分组
        by_sev_threats: dict[str, list[dict]] = {}
        for t in threats:
            by_sev_threats.setdefault(t["severity"], []).append(t)
        for sev in sorted(by_sev_threats, key=_sev):
            items = by_sev_threats[sev]
            lines.append(f"### {sev}（{len(items)} 条）")
            lines.append("")
            for i, t in enumerate(items, 1):
                lines.append(f"{i}. **{t['title']}**（{t['type']}）")
                if t.get("cwe"):
                    lines.append(f"   - CWE：{t['cwe']}")
                dread = t.get("dread")
                if isinstance(dread, dict):
                    ds = sum(dread.values())
                    lines.append(f"   - DREAD：{ds}/50（危害:{dread.get('damage',0)} 可重复:{dread.get('reproducibility',0)} 可利用:{dread.get('exploitability',0)} 影响:{dread.get('affectedUsers',0)} 可发现:{dread.get('discoverability',0)}）")
                lines.append(f"   - 组件：{t['component'] or '-'}")
                lines.append(f"   - 状态：{t['status']}")
                lines.append(f"   - 描述：{t['description'] or '-'}")
                lines.append(f"   - 缓解措施：{t['mitigation'] or '-'}")
                refs = t.get("references") or []
                if refs:
                    lines.append(f"   - 参考：{'；'.join(refs)}")
                lines.append("")
    return "\n".join(lines)


def _pct(rate) -> str:
    """把 0~1 比率格式化为百分比字符串。"""
    try:
        return f"{round(float(rate) * 100, 1)}%"
    except (TypeError, ValueError):
        return "-"


def _collect_threats(model: dict[str, Any]) -> list[dict[str, Any]]:
    """从 Threat Dragon 模型中收集所有威胁，并补充所属组件名。"""
    out: list[dict[str, Any]] = []
    diagrams = ((model.get("detail") or {}).get("diagrams")) or []
    for diagram in diagrams:
        cells = diagram.get("cells") or []
        name_by_cell = {
            c["id"]: ((c.get("data") or {}).get("name") or "")
            for c in cells
        }
        for cell in cells:
            for threat in cell.get("threats") or []:
                t = dict(threat)
                t["component"] = name_by_cell.get(cell["id"], "")
                out.append(t)
    return out


def _to_td_json(record: dict[str, Any]) -> dict[str, Any]:
    """返回 Threat Dragon v2 标准模型 JSON（可直接导入官方 TD）。"""
    model = record.get("model") or {}
    td = dict(model)
    detail = td.get("detail") or {}
    if not detail.get("version"):
        detail = dict(detail)
        detail["version"] = "0.1.0"
    td["detail"] = detail
    td.setdefault("summary", {})
    td["summary"]["title"] = record.get("title", td["summary"].get("title", "Threat Model"))
    return td


def render_result_json(record: dict[str, Any]) -> str:
    """把结果渲染为 Threat Dragon 标准 JSON 字符串。"""
    return json.dumps(_to_td_json(record), ensure_ascii=False, indent=2)


def _dread_total(dread) -> int:
    if isinstance(dread, dict):
        return int(dread.get("total", sum(dread.values())))
    return 0


def render_result_csv(record: dict[str, Any]) -> str:
    """把威胁清单渲染为 CSV 文本（UTF-8，含 BOM 以便 Excel 打开中文）。"""
    model = record.get("model") or {}
    threats = _collect_threats(model)
    fieldnames = [
        "severity", "type", "title", "component", "status",
        "cwe", "dread_total", "description", "mitigation", "references",
    ]
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=fieldnames)
    writer.writeheader()
    for t in threats:
        refs = t.get("references") or []
        writer.writerow({
            "severity": t.get("severity", ""),
            "type": t.get("type", ""),
            "title": t.get("title", ""),
            "component": t.get("component", ""),
            "status": t.get("status", ""),
            "cwe": t.get("cwe", ""),
            "dread_total": _dread_total(t.get("dread")),
            "description": t.get("description", ""),
            "mitigation": t.get("mitigation", ""),
            "references": "；".join(refs),
        })
    return "\ufeff" + buf.getvalue()


def render_result_docx(record: dict[str, Any]) -> bytes:
    """把一条结果记录渲染为 Word (.docx) 报告（二进制字节）。

    依赖 python-docx；若服务器未安装则抛出 RuntimeError，由上层转为 5xx。
    """
    from docx import Document  # type: ignore
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    model = record.get("model") or {}
    summary = model.get("summary") or {}
    stats = record.get("stats") or {}
    methodology = record.get("methodology") or "STRIDE"
    created = _fmt_time(record.get("created_at", 0))

    doc = Document()

    # 中文字体（宋体），确保中文正常显示
    def _set_font(style, name="宋体", size=None, bold=None, color=None):
        try:
            style.font.name = name
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), name)
        except Exception:
            pass
        if size is not None:
            style.font.size = Pt(size)
        if bold is not None:
            style.font.bold = bold
        if color is not None:
            style.font.color.rgb = color

    _set_font(doc.styles["Normal"], size=11)
    _set_font(doc.styles["Heading 1"], size=18, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))
    _set_font(doc.styles["Heading 2"], size=15, bold=True, color=RGBColor(0x2B, 0x57, 0x9A))
    _set_font(doc.styles["Heading 3"], size=13, bold=True, color=RGBColor(0x2B, 0x57, 0x9A))

    # 标题
    title = doc.add_heading(record.get("title", "威胁建模报告"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"创建时间：{created}    方法论：{methodology}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    def _kv(label: str, value):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}：")
        r1.bold = True
        p.add_run(str(value) if value else "-")

    # 1、统计概览
    doc.add_heading("1、统计概览", level=2)
    _kv("组件数", stats.get("componentCount", "-"))
    _kv("数据流数", stats.get("flowCount", "-"))
    _kv("威胁总数", stats.get("threatCount", "-"))
    by_sev = stats.get("threatCountBySeverity") or {}
    if by_sev:
        bits = [f"{k}: {v}" for k, v in sorted(
            by_sev.items(), key=lambda kv: _sev(kv[0])
        )]
        _kv("威胁按等级", "，".join(bits))
    industry = stats.get("industry")
    if industry:
        try:
            from .ai_knowledge import get_industry_template
            tmpl = get_industry_template(industry)
            _kv("行业场景", tmpl["label"] if tmpl else industry)
        except Exception:
            _kv("行业场景", industry)
    _kv("模型标题", summary.get("title", "-"))
    _kv("模型描述", summary.get("description", "-") or "-")

    # 度量指标
    metrics = stats.get("metrics") or {}
    if metrics:
        doc.add_heading("1.2、度量指标", level=2)
        _kv("元素覆盖度", f"{_pct(metrics.get('coverageRate'))}（{metrics.get('modeledElements', 0)}/{metrics.get('totalElements', 0)}）")
        _kv("高风险威胁收敛率", _pct(metrics.get("riskConvergence")))
        dread_avg = metrics.get("dreadAverage")
        if dread_avg:
            dread_label = {
                "damage": "危害", "reproducibility": "可重复性",
                "exploitability": "可利用性", "affectedUsers": "受影响面",
                "discoverability": "可发现性",
            }
            bits = [f"{dread_label.get(k, k)}:{v}" for k, v in dread_avg.items() if k != "total"]
            _kv("DREAD 均值", " / ".join(bits) + f"（综合 {dread_avg.get('total', 0)}）")
        llm_cov = metrics.get("owaspLlmCoverRate")
        if llm_cov is not None:
            covered = metrics.get("owaspLlmCovered") or []
            _kv("OWASP Top10 for LLM 覆盖", f"{_pct(llm_cov)}（{', '.join(covered) if covered else '无'}）")
        compliance = metrics.get("compliance")
        if compliance:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("合规映射：").bold = True
            for item in compliance:
                mark = "✓" if item.get("covered") else "✗"
                doc.add_paragraph(f"  {mark} {item.get('code', '')} {item.get('label', '')}", style="List Bullet")

    # 2、威胁明细
    threats = _collect_threats(model)
    doc.add_heading(f"2、威胁明细（共 {len(threats)} 条）", level=2)

    if not threats:
        doc.add_paragraph("未识别到威胁。")
    else:
        by_sev_threats: dict[str, list[dict]] = {}
        for t in threats:
            by_sev_threats.setdefault(t["severity"], []).append(t)
        for sev in sorted(by_sev_threats, key=_sev):
            items = by_sev_threats[sev]
            doc.add_heading(f"{sev}（{len(items)} 条）", level=3)
            table = doc.add_table(rows=1, cols=6)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for i, h in enumerate(["序号", "威胁", "类型", "组件", "CWE", "DREAD"]):
                hdr[i].text = h
                for run in hdr[i].paragraphs[0].runs:
                    run.bold = True
            for i, t in enumerate(items, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = t.get("title", "")
                row[2].text = t.get("type", "")
                row[3].text = t.get("component", "")
                row[4].text = t.get("cwe", "")
                row[5].text = str(_dread_total(t.get("dread")))
            # 每条威胁的详情描述
            for i, t in enumerate(items, 1):
                doc.add_paragraph()
                p = doc.add_paragraph()
                r = p.add_run(f"{i}. {t.get('title', '')}（{t.get('type', '')}）")
                r.bold = True
                if t.get("cwe"):
                    doc.add_paragraph(f"CWE：{t['cwe']}")
                dread = t.get("dread")
                if isinstance(dread, dict):
                    doc.add_paragraph(
                        f"DREAD：{_dread_total(dread)}/50"
                        f"（危害:{dread.get('damage',0)} 可重复:{dread.get('reproducibility',0)}"
                        f" 可利用:{dread.get('exploitability',0)} 影响:{dread.get('affectedUsers',0)}"
                        f" 可发现:{dread.get('discoverability',0)}）"
                    )
                doc.add_paragraph(f"状态：{t.get('status', '')}")
                doc.add_paragraph(f"描述：{t.get('description', '') or '-'}")
                doc.add_paragraph(f"缓解措施：{t.get('mitigation', '') or '-'}")
                refs = t.get("references") or []
                if refs:
                    doc.add_paragraph(f"参考：{'；'.join(refs)}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
