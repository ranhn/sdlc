"""文档内容抽取服务：支持 .txt / .md / .pdf / .docx。

用于前端「上传文档 → AI 自动解析」：仅抽取纯文本返回，
不调用 LLM，抽取结果直接填充到对应输入框，由用户决定后续是否建模。
"""

from __future__ import annotations

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# 最大抽取字符数：防止超大文档导致前端/LLM 压力过大
MAX_EXTRACT_CHARS = 60000


class UnsupportedFileTypeError(ValueError):
    """不支持的文件类型。"""


def _friendly_parse_error(ext: str, exc: BaseException) -> str:
    """把底层 PDF/DOCX 库异常翻译成用户能看懂的提示。

    设计目标：不要把 "文档解析失败" 这种无用信息直接抛给前端。
    常见原因（按概率排）：加密 PDF / PDF 截断 / PDF 格式损坏 / DOCX 损坏 /
    扫描版 PDF（pypdf 抽不出文字时不会抛异常，会返回空字符串，那种由调用方
    单独检查 "extracted 文本为空" 来报）。
    """
    name = type(exc).__name__
    msg = str(exc).strip()
    e_lower = (msg or name).lower()
    if ext == ".pdf":
        if "encrypted" in e_lower or "decrypted" in e_lower or "password" in e_lower:
            return "PDF 已加密，无法解析（请先用 PDF 工具去除密码后再上传）"
        if name in ("EofError",) or "unexpected eof" in e_lower:
            return "PDF 文件不完整（被截断），请重新下载/导出后上传"
        if "xref" in e_lower or "startxref" in e_lower or "invalid" in e_lower:
            return f"PDF 格式损坏或不规范（{msg or name}），请用 PDF 工具重新导出后再上传"
    if ext == ".docx":
        if "not a zip file" in e_lower or "bad zipfile" in e_lower:
            return "DOCX 文件格式损坏（不是有效的 zip 容器），请用 Word 重新保存后上传"
    # 兜底：把底层错误透出
    return f"文档解析失败（{ext}）：{msg or name}"

def extract_text(filename: str, data: bytes) -> str:
    """根据文件扩展名抽取文本。

    Args:
        filename: 原始文件名（带扩展名）。
        data: 文件二进制内容。

    Returns:
        抽取出的纯文本（已做裁剪与清理）。

    Raises:
        UnsupportedFileTypeError: 扩展名不受支持。
    """
    ext = _extension(filename)
    try:
        if ext in (".txt", ".md", ".markdown", ".text"):
            text = _extract_plain(data)
        elif ext == ".pdf":
            text = _extract_pdf(data)
        elif ext == ".docx":
            text = _extract_docx(data)
        else:
            raise UnsupportedFileTypeError(
                f"不支持的文件类型「{ext or '(无扩展名)'}」，仅支持 .txt / .md / .pdf / .docx"
            )
    except UnsupportedFileTypeError:
        raise
    except Exception as e:  # noqa: BLE001 —— 解析失败统一转成用户可读错误
        logger.warning("文档解析失败 %s: %s", filename, e, exc_info=True)
        raise ValueError(_friendly_parse_error(ext, e)) from e

    return _clean(text)


def extract_assets(filename: str, data: bytes) -> dict:
    """抽取文档的文本 + 内嵌图片（data URI 列表）。

    Args:
        filename: 原始文件名（带扩展名）。
        data: 文件二进制内容。

    Returns:
        {
            "text": str,
            "images": [data_uri, ...],
            "warnings": [str, ...],   # P1-5：抽取过程中的非阻塞告警（Pillow 缺失等）
        }
        images 仅对含内嵌图片的 PDF / DOCX 返回；纯文本文件 images 为空列表。
    """
    ext = _extension(filename)
    images: list[str] = []
    text = ""
    warnings: list[str] = []   # P1-5

    # P1-5：检查 Pillow 是否可用；PDF 图片抽取强依赖它
    if ext == ".pdf":
        try:
            from PIL import Image  # noqa: F401
        except ImportError:
            warnings.append(
                "服务器未安装 Pillow 库，PDF 中的架构图将无法抽取（仅返回文本）。"
                "请在 backend/requirements.txt 中加上 Pillow 后重启服务。"
            )

    try:
        if ext in (".txt", ".md", ".markdown", ".text"):
            text = _extract_plain(data)
        elif ext == ".pdf":
            text, images = _extract_pdf_assets(data)
            # P1-5：如果原本该文档含图但一张都没抽出来，给个提示
            if not images and _pdf_likely_has_images(data):
                warnings.append(
                    "PDF 中可能含图但未能抽取（可能是扫描版 PDF 或依赖库未安装 Pillow）。"
                )
        elif ext == ".docx":
            text, images = _extract_docx_assets(data)
        else:
            raise UnsupportedFileTypeError(
                f"不支持的文件类型「{ext or '(无扩展名)'}」，仅支持 .txt / .md / .pdf / .docx"
            )
    except UnsupportedFileTypeError:
        raise
    except Exception as e:  # noqa: BLE001
        logger.warning("文档解析失败 %s: %s", filename, e, exc_info=True)
        raise ValueError(_friendly_parse_error(ext, e)) from e

    return {
        "text": _clean(text),
        "images": images[:24],  # P1-6：上限从 12 提到 24
        "warnings": warnings,
    }


def _pdf_likely_has_images(data: bytes) -> bool:
    """启发式判断 PDF 是否含图（避免每次都解包整个文件做精确检查）。"""
    if not data:
        return False
    # 关键标记：PDF 对象流里含 /Subtype /Image（按 KB 量级粗筛，避免全文件扫描）
    head = data[:2 * 1024 * 1024]   # 头 2MB 足够
    return b"/Subtype" in head and b"/Image" in head


def _extension(filename: str) -> str:
    name = (filename or "").strip().lower()
    if "." not in name:
        return ""
    return name[name.rfind("."):]


def _extract_plain(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    # 兜底
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("服务器未安装 pypdf，无法解析 PDF") from e

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 —— 单页失败不影响整体
            continue
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("服务器未安装 python-docx，无法解析 Word 文档") from e

    doc = Document(io.BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    # 表格内容一并抽取
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(c for c in cells if c))
    return "\n".join(lines)


def _extract_pdf_assets(data: bytes) -> tuple[str, list[str]]:
    """抽取 PDF 文本与内嵌图片（data URI 列表）。"""
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("服务器未安装 pypdf，无法解析 PDF") from e

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    images: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            continue
        try:
            for img in getattr(page, "images", []) or []:
                try:
                    uri = _image_to_data_uri(img.data, img.name)
                    if uri:
                        images.append(uri)
                except Exception as exc:  # noqa: BLE001
                    # 不再静默吞掉：pypdf.page.images 需要 Pillow，否则会 ImportError 导致 0 张架构图
                    logger.warning("page image skipped: %s: %s", type(exc).__name__, exc)
                    continue
        except Exception as exc:  # noqa: BLE001
            logger.warning("page.images traversal failed: %s: %s", type(exc).__name__, exc)
            continue
    return "\n".join(parts), images


def _extract_docx_assets(data: bytes) -> tuple[str, list[str]]:
    """抽取 DOCX 文本与内嵌图片（data URI 列表）。

    P1-4：补充抽取"浮动图"（inline_shapes）—— 之前只抽 word/media/*，
    但用户拖拽/Word 插入的浮动图（drawings）有时只有 image 引用而没有
    独立 word/media/ 文件（取决于 Word 版本/图片大小/插入方式）。
    使用 ``doc.inline_shapes`` + ``_get_image_bytes`` 获取其字节流。
    """
    import zipfile

    parts: list[str] = []
    images: list[str] = []

    # 先抽取文本（保持与 extract_text 一致）
    try:
        from docx import Document  # type: ignore
    except ImportError as e:  # pragma: no cover
        raise RuntimeError("服务器未安装 python-docx，无法解析 Word 文档") from e

    try:
        doc = Document(io.BytesIO(data))
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(c for c in cells if c))

        # P1-4：抽取段落中的内嵌浮动图（drawings/inline_shapes）
        # 这类图的 binary 不一定落在 word/media/*，需通过 image_part 直接拿
        for shape_idx, shape in enumerate(doc.inline_shapes or []):
            try:
                image_part = getattr(shape, "image", None)
                if image_part is None:
                    continue
                # 新版 python-docx 的 InlineShape.image 返回 ImagePart，
                # 其 .blob 是图片二进制；旧版可能为 None
                blob = getattr(image_part, "blob", None)
                if not blob:
                    continue
                # 扩展名用 image_part.content_type 推断
                mime = getattr(image_part, "content_type", "image/png") or "image/png"
                ext = {
                    "image/png": ".png",
                    "image/jpeg": ".jpg",
                    "image/gif": ".gif",
                    "image/bmp": ".bmp",
                    "image/webp": ".webp",
                }.get(mime.split(";")[0].strip(), ".png")
                uri = _image_to_data_uri(blob, f"inline_{shape_idx}{ext}")
                if uri:
                    images.append(uri)
            except Exception as exc:  # noqa: BLE001
                logger.warning("DOCX inline shape %d 抽取失败: %s", shape_idx, exc)
                continue
    except Exception as e:  # noqa: BLE001
        logger.warning("DOCX 文本/浮动图抽取失败，仅尝试 word/media/*: %s", e)

    # 抽取内嵌图片（word/media/*）
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for name in zf.namelist():
                low = name.lower()
                if not low.startswith("word/media/") or not _looks_image(low):
                    continue
                try:
                    raw = zf.read(name)
                    uri = _image_to_data_uri(raw, name)
                    if uri:
                        images.append(uri)
                except Exception:  # noqa: BLE001
                    continue
    except Exception as e:  # noqa: BLE001
        logger.warning("DOCX 图片抽取失败: %s", e)

    # 去重（同一张图可能既在 word/media/ 又在 inline_shapes 里被引用）
    if images:
        seen: set[str] = set()
        deduped: list[str] = []
        for u in images:
            if u in seen:
                continue
            seen.add(u)
            deduped.append(u)
        if len(deduped) != len(images):
            logger.info("DOCX 图片去重：%d → %d", len(images), len(deduped))
        images = deduped

    return "\n".join(parts), images


def _looks_image(name: str) -> bool:
    return name.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"))


# 图片转 data URI 前的压缩参数：防止超大 base64 拖垮前端/LLM
_IMAGE_MAX_DIM = 1200        # 最长边像素上限
_IMAGE_JPEG_QUALITY = 75     # JPEG 压缩质量
# 每张图片 data URI 的字节数上限；超过则进一步降采样
_IMAGE_MAX_BYTES = 2_200_000  # ~2.2MB（含 base64 膨胀，原始字节约 1.6MB）


def _normalize_image(raw: bytes) -> bytes:
    """对图片做降采样 / 压缩，返回可安全内嵌的字节。

    优先使用 Pillow 高质量压缩；未安装时若图片超过上限则直接丢弃
    （由调用方忽略），避免超大图拖垮前端。不能识别 / 处理失败时
    返回原字节（交由上层判断是否超限）。
    """
    if not raw:
        return raw
    try:
        from PIL import Image  # type: ignore
    except ImportError:
        # 无 Pillow：若原始字节已超上限，返回空以丢弃；否则原样返回
        return b"" if len(raw) > _IMAGE_MAX_BYTES else raw

    try:
        with Image.open(io.BytesIO(raw)) as img:
            img.load()
            # 统一转 RGB，避免透明 PNG 压缩成 JPEG 后出现黑底
            rgb = img.convert("RGB") if img.mode in ("RGBA", "P", "LA") else img
            if max(rgb.size) > _IMAGE_MAX_DIM:
                rgb.thumbnail((_IMAGE_MAX_DIM, _IMAGE_MAX_DIM), Image.LANCZOS)
            buf = io.BytesIO()
            # 依据预估体积决定是否转 JPEG：原始就是小 PNG 则保留 PNG
            if len(raw) <= _IMAGE_MAX_BYTES and rgb.format == "PNG":
                rgb.save(buf, format="PNG", optimize=True)
            else:
                rgb.save(buf, format="JPEG", quality=_IMAGE_JPEG_QUALITY, optimize=True)
            out = buf.getvalue()
            return out if len(out) <= _IMAGE_MAX_BYTES else b""
    except Exception:  # noqa: BLE001 —— 解码失败则走原始字节判定
        return raw


def _image_to_data_uri(raw: bytes, name: str) -> Optional[str]:
    """把图片字节转成 data URI（内嵌前先降采样压缩）。不能识别类型时返回 None。"""
    import base64
    import mimetypes

    if not raw:
        return None
    mime, _ = mimetypes.guess_type(name)
    if not mime or not mime.startswith("image/"):
        # 由扩展名推断
        ext = _extension(name)
        mime = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
            ".webp": "image/webp",
        }.get(ext)
        if not mime:
            return None
    raw = _normalize_image(raw)
    if not raw:
        return None
    b64 = base64.b64encode(raw).decode("ascii")
    return f"data:{mime};base64,{b64}"


def _clean(text: str) -> str:
    """裁剪并清理抽取文本。"""
    if not text:
        return ""
    # 合并 3 个以上连续换行为两个换行（压缩多余空行）
    text = _collapse_blank_lines(text)
    # 限制长度
    if len(text) > MAX_EXTRACT_CHARS:
        text = text[:MAX_EXTRACT_CHARS] + "\n\n[内容过长，已截断…]"
    return text


def _collapse_blank_lines(text: str) -> str:
    import re

    return re.sub(r"\n{3,}", "\n\n", text)


def allowed_extensions() -> list[str]:
    return [".txt", ".md", ".markdown", ".text", ".pdf", ".docx"]
