"""漏洞管理路由：提交/确认/修复/复测/关闭 + 状态机 + 评论。"""
import csv
import io
import json

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import inspect
from sqlalchemy.orm import Session

from ..database import get_db
from ..models import User, Vuln, VulnComment, VulnFlow
from ..schemas import (
    VulnAssign,
    VulnCommentOut,
    VulnCreate,
    VulnFlowOut,
    VulnOut,
    VulnReject,
    VulnStatusAction,
)
from ..security import get_current_user, write_operation_log
from ..state_machine import TRANSITIONS, validate_action

router = APIRouter(prefix="/api/vulns", tags=["漏洞管理"])

STATUS_NAMES = {
    "draft": "草稿", "pending": "待确认", "confirmed": "已确认", "fixing": "修复中",
    "retest": "待复测", "fixed": "已修复", "closed": "已关闭", "rejected": "已驳回", "ignored": "已忽略",
}


def _to_out(v: Vuln, db: Session) -> VulnOut:
    data = {c.key: getattr(v, c.key) for c in inspect(v).mapper.column_attrs}
    if data.get("screenshots"):
        try:
            data["screenshots"] = json.loads(data["screenshots"])
        except (TypeError, json.JSONDecodeError):
            data["screenshots"] = None
    else:
        data["screenshots"] = None
    if data.get("step_screenshots"):
        try:
            data["step_screenshots"] = json.loads(data["step_screenshots"])
        except (TypeError, json.JSONDecodeError):
            data["step_screenshots"] = None
    else:
        data["step_screenshots"] = None
    out = VulnOut.model_validate(data)
    reporter = db.query(User).filter(User.id == v.reporter_id).first()
    assignee = db.query(User).filter(User.id == v.assignee_id).first() if v.assignee_id is not None else None
    reviewer = db.query(User).filter(User.id == v.reviewer_id).first() if v.reviewer_id is not None else None
    out.reporter_name = reporter.full_name if reporter else None
    out.assignee_name = assignee.full_name if assignee else None
    out.reviewer_name = reviewer.full_name if reviewer else None
    if v.system:
        out.system_name = v.system.name
    return out


def _record_flow(db: Session, vuln_id: int, from_status: str | None, to_status: str,
                 operator: User, comment: str | None):
    db.add(VulnFlow(
        vuln_id=vuln_id,
        from_status=from_status,
        to_status=to_status,
        operator_id=operator.id,
        operator_name=operator.full_name,
        comment=comment,
    ))
    db.commit()


@router.get("", response_model=list[VulnOut])
def list_vulns(
    status: str | None = Query(default=None),
    severity: str | None = Query(default=None),
    system_id: int | None = Query(default=None),
    mine: bool = Query(default=False),
    assigned_to_me: bool = Query(default=False),
    is_external: bool | None = Query(default=None, description="漏洞来源过滤：None=全部, True=外部, False=内部"),
    db: Session = Depends(get_db),
    current: User = Depends(get_current_user),
):
    query = db.query(Vuln)
    if status:
        query = query.filter(Vuln.status == status)
    if severity:
        query = query.filter(Vuln.severity == severity)
    if system_id:
        query = query.filter(Vuln.system_id == system_id)
    if is_external is not None:
        query = query.filter(Vuln.is_external == is_external)
    if mine:
        query = query.filter((Vuln.reporter_id == current.id) | (Vuln.assignee_id == current.id))
    if assigned_to_me:
        query = query.filter(Vuln.assignee_id == current.id)
    vulns = query.order_by(Vuln.created_at.desc()).all()
    return [_to_out(v, db) for v in vulns]


@router.get("/export")
def export_vulns(
    fmt: str = Query(..., pattern="^(csv|docx)$"),
    status: str | None = Query(default=None),
    severity: str | None = Query(default=None),
    system_id: int | None = Query(default=None),
    mine: bool = Query(default=False),
    assigned_to_me: bool = Query(default=False),
    db: Session = Depends(get_db),
    current: User = Depends(get_current_user),
):
    """批量导出漏洞（CSV / Word）。仅管理员/安全专家可操作。"""
    if current.role is None or current.role.code not in ("admin", "secops"):
        raise HTTPException(status_code=403, detail="仅管理员/安全专家可导出漏洞")
    query = db.query(Vuln)
    if status:
        query = query.filter(Vuln.status == status)
    if severity:
        query = query.filter(Vuln.severity == severity)
    if system_id:
        query = query.filter(Vuln.system_id == system_id)
    if mine:
        query = query.filter((Vuln.reporter_id == current.id) | (Vuln.assignee_id == current.id))
    if assigned_to_me:
        query = query.filter(Vuln.assignee_id == current.id)
    vulns = query.order_by(Vuln.created_at.desc()).all()
    rows = [_to_out(v, db) for v in vulns]

    if fmt == "csv":
        return _export_csv(rows)
    return _export_docx(rows)


def _export_csv(rows: list[VulnOut]):
    headers = ["ID", "标题", "所属系统", "等级", "类型", "状态", "提交人", "负责人", "复测人", "创建时间"]
    buf = io.StringIO()
    # 写入 BOM 让 Excel 正确识别 UTF-8
    buf.write("\ufeff")
    writer = csv.writer(buf)
    writer.writerow(headers)
    for r in rows:
        writer.writerow([
            r.id, r.title, r.system_name or "",
            {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危"}.get(r.severity, r.severity),
            r.vuln_type or "", STATUS_NAMES.get(r.status, r.status),
            r.reporter_name or "", r.assignee_name or "未指派", r.reviewer_name or "",
            r.created_at.strftime("%Y-%m-%d %H:%M") if r.created_at else "",
        ])
    data = buf.getvalue().encode("utf-8")
    return StreamingResponse(
        iter([data]),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": 'attachment; filename="vulns.csv"'},
    )


def _export_docx(rows: list[VulnOut]):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    doc.add_heading("漏洞清单", level=1)
    doc.add_paragraph(f"导出时间：{__import__('datetime').datetime.now().strftime('%Y-%m-%d %H:%M')}    共 {len(rows)} 条")
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["ID", "标题", "系统", "等级", "类型", "状态", "负责人", "创建时间"]):
        hdr[i].text = h
    sev_map = {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危"}
    for r in rows:
        cells = table.add_row().cells
        cells[0].text = str(r.id)
        cells[1].text = r.title or ""
        cells[2].text = r.system_name or ""
        cells[3].text = sev_map.get(r.severity, r.severity or "")
        cells[4].text = r.vuln_type or ""
        cells[5].text = STATUS_NAMES.get(r.status, r.status or "")
        cells[6].text = r.assignee_name or "未指派"
        cells[7].text = r.created_at.strftime("%Y-%m-%d %H:%M") if r.created_at else ""
    # 详情段落
    if rows:
        doc.add_paragraph()
        doc.add_heading("漏洞详情", level=2)
        for r in rows:
            doc.add_heading(f"#{r.id} {r.title}", level=3)
            doc.add_paragraph(f"所属系统：{r.system_name or '—'}    等级：{sev_map.get(r.severity, r.severity or '')}    状态：{STATUS_NAMES.get(r.status, r.status or '')}")
            doc.add_paragraph(f"提交人：{r.reporter_name or '—'}    负责人：{r.assignee_name or '未指派'}    复测人：{r.reviewer_name or '—'}")
            if r.description:
                doc.add_paragraph(f"【漏洞描述】{r.description}")
            if r.reproduce_steps:
                doc.add_paragraph(f"【复现步骤】{r.reproduce_steps}")
            if r.impact:
                doc.add_paragraph(f"【影响范围】{r.impact}")
            if r.step_screenshots:
                doc.add_paragraph(f"【步骤截图】共 {len(r.step_screenshots)} 张（图片需通过系统查看）")
            doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="vulns.docx"'},
    )


@router.get("/{vuln_id}", response_model=VulnOut)
def get_vuln(vuln_id: int, db: Session = Depends(get_db), current: User = Depends(get_current_user)):
    v = db.query(Vuln).filter(Vuln.id == vuln_id).first()
    if not v:
        raise HTTPException(status_code=404, detail="漏洞不存在")
    return _to_out(v, db)


@router.post("", response_model=VulnOut, status_code=201)
def create_vuln(data: VulnCreate, db: Session = Depends(get_db), current: User = Depends(get_current_user)):
    v = Vuln(
        title=data.title,
        description=data.description,
        reproduce_steps=data.reproduce_steps,
        impact=data.impact,
        screenshots=json.dumps(data.screenshots, ensure_ascii=False) if data.screenshots else None,
        step_screenshots=json.dumps(data.step_screenshots, ensure_ascii=False) if data.step_screenshots else None,
        system_id=data.system_id,
        severity=data.severity,
        vuln_type=data.vuln_type,
        cvss=data.cvss,
        reporter_id=current.id,
        assignee_id=data.assignee_id,
        status="pending",
        source="manual",
        is_external=data.is_external,
        external_source=data.external_source,
    )
    db.add(v)
    db.commit()
    db.refresh(v)
    _record_flow(db, v.id, "draft", "pending", current, "漏洞提交")
    write_operation_log(db, current, "create_vuln", "vuln", f"提交漏洞 #{v.id} {v.title}")
    return _to_out(v, db)


@router.post("/{vuln_id}/assign", response_model=VulnOut)
def assign_vuln(vuln_id: int, data: VulnAssign, db: Session = Depends(get_db),
                current: User = Depends(get_current_user)):
    v = db.query(Vuln).filter(Vuln.id == vuln_id).first()
    if not v:
        raise HTTPException(status_code=404, detail="漏洞不存在")
    if current.role is None or current.role.code not in ("admin", "secops"):
        raise HTTPException(status_code=403, detail="仅安全专家可指派")
    if not db.query(User).filter(User.id == data.assignee_id).first():
        raise HTTPException(status_code=400, detail="负责人不存在")
    assignee = db.query(User).filter(User.id == data.assignee_id).first()
    v.assignee_id = data.assignee_id
    db.commit()
    db.refresh(v)
    # 详情里同时记录「指派给 谁(中文名)[ID]」，便于审计日志辨识接收人
    assignee_desc = (
        f"{assignee.username}({assignee.full_name})[ID={assignee.id}]"
        if assignee
        else f"ID={data.assignee_id}"
    )
    write_operation_log(
        db, current, "assign_vuln", "vuln",
        f"指派漏洞 #{v.id}「{v.title}」给 {assignee_desc}",
    )
    return _to_out(v, db)


@router.post("/{vuln_id}/action/{action}", response_model=VulnOut)
def vuln_action(vuln_id: int, action: str, data: VulnStatusAction,
                db: Session = Depends(get_db), current: User = Depends(get_current_user)):
    v = db.query(Vuln).filter(Vuln.id == vuln_id).first()
    if not v:
        raise HTTPException(status_code=404, detail="漏洞不存在")
    role_code = current.role.code if current.role else "user"
    if not validate_action(action, v.status, role_code):
        raise HTTPException(status_code=403, detail=f"当前状态({STATUS_NAMES.get(v.status, v.status)})下，角色无权执行[{action}]操作")

    to_status = TRANSITIONS[action]
    from_status = v.status
    v.status = to_status

    if action == "close":
        v.closed_at = __import__("datetime").datetime.utcnow()
    if action == "finish_fix":
        v.fixed_at = __import__("datetime").datetime.utcnow()
        v.reviewer_id = current.id
    if action == "pass_retest":
        v.reviewer_id = current.id

    db.commit()
    db.refresh(v)
    _record_flow(db, v.id, from_status, to_status, current, data.comment)
    write_operation_log(db, current, f"vuln_{action}", "vuln", f"漏洞 #{v.id} {from_status}->{to_status}")
    return _to_out(v, db)


@router.post("/{vuln_id}/reject", response_model=VulnOut)
def reject_vuln(vuln_id: int, data: VulnReject, db: Session = Depends(get_db),
                current: User = Depends(get_current_user)):
    v = db.query(Vuln).filter(Vuln.id == vuln_id).first()
    if not v:
        raise HTTPException(status_code=404, detail="漏洞不存在")
    role_code = current.role.code if current.role else "user"
    if not validate_action("reject", v.status, role_code):
        raise HTTPException(status_code=403, detail="无权限驳回")
    v.status = "rejected"
    v.rejection_reason = data.reason
    db.commit()
    db.refresh(v)
    _record_flow(db, v.id, "pending", "rejected", current, f"驳回：{data.reason}")
    write_operation_log(db, current, "vuln_reject", "vuln", f"漏洞 #{v.id} 驳回：{data.reason}")
    return _to_out(v, db)


@router.get("/{vuln_id}/flows", response_model=list[VulnFlowOut])
def vuln_flows(vuln_id: int, db: Session = Depends(get_db), current: User = Depends(get_current_user)):
    return db.query(VulnFlow).filter(VulnFlow.vuln_id == vuln_id).order_by(VulnFlow.created_at.asc()).all()


@router.get("/{vuln_id}/comments", response_model=list[VulnCommentOut])
def vuln_comments(vuln_id: int, db: Session = Depends(get_db), current: User = Depends(get_current_user)):
    return db.query(VulnComment).filter(VulnComment.vuln_id == vuln_id).order_by(VulnComment.created_at.asc()).all()


@router.post("/{vuln_id}/comments", response_model=VulnCommentOut)
def add_comment(vuln_id: int, content: VulnStatusAction, db: Session = Depends(get_db),
                current: User = Depends(get_current_user)):
    if not db.query(Vuln).filter(Vuln.id == vuln_id).first():
        raise HTTPException(status_code=404, detail="漏洞不存在")
    c = VulnComment(vuln_id=vuln_id, user_id=current.id, username=current.full_name, content=content.comment)
    db.add(c)
    db.commit()
    db.refresh(c)
    return c


@router.delete("/{vuln_id}")
def delete_vuln(vuln_id: int, db: Session = Depends(get_db), current: User = Depends(get_current_user)):
    """删除漏洞。仅管理员/安全专家可操作。"""
    if current.role is None or current.role.code not in ("admin", "secops"):
        raise HTTPException(status_code=403, detail="仅管理员/安全专家可删除漏洞")
    v = db.query(Vuln).filter(Vuln.id == vuln_id).first()
    if not v:
        raise HTTPException(status_code=404, detail="漏洞不存在")
    title = v.title
    # 级联删除关联数据
    db.query(VulnFlow).filter(VulnFlow.vuln_id == vuln_id).delete()
    db.query(VulnComment).filter(VulnComment.vuln_id == vuln_id).delete()
    db.delete(v)
    db.commit()
    write_operation_log(db, current, "delete_vuln", "vuln", f"删除漏洞 #{vuln_id} {title}")
    return {"detail": "已删除"}
